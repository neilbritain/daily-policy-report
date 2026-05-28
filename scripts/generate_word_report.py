#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 报告生成脚本 v2.0
从 reports/latest.json 读取数据，生成格式化 Word 文档
"""
import json
from pathlib import Path
from datetime import datetime, timedelta, timezone

BEIJING_TZ = timezone(timedelta(hours=8))


def get_report_data() -> dict:
    """优先读 latest.json，找不到则报错"""
    json_path = Path("reports/latest.json")
    if json_path.exists():
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)
    raise FileNotFoundError(
        "reports/latest.json 不存在，请先运行 generate_report.py"
    )


def save_word_report(report_data: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    report_dir = Path("reports")
    report_dir.mkdir(exist_ok=True)
    date = report_data["date"]

    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # 默认字体
    doc.styles["Normal"].font.name = "仿宋"
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    # ── 标题 ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = title_p.add_run(report_data["title"])
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(91, 109, 228)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_r = sub_p.add_run(f"生成日期：{date}　　数据 · 算力 · 信息化")
    sub_r.font.size = Pt(11)
    sub_r.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_paragraph()  # 空行

    # ── 正文节 ──
    def add_section(heading: str, policies: list):
        # 节标题
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(91, 109, 228)
        h.runs[0].font.size = Pt(14)

        for p in policies:
            # 政策标题
            title_para = doc.add_paragraph()
            tr = title_para.add_run(f"《{p['title']}》")
            tr.font.bold = True
            tr.font.size = Pt(12)

            # 元信息
            meta_para = doc.add_paragraph()
            meta_text = (
                f"文　　号：{p.get('doc_number', '暂无')}\n"
                f"发文机关：{p['institution']}\n"
                f"发文日期：{p['date']}"
            )
            if p.get("url"):
                meta_text += f"\n原文链接：{p['url']}"
            mr = meta_para.add_run(meta_text)
            mr.font.size = Pt(10)
            mr.font.color.rgb = RGBColor(100, 100, 100)

            # 摘要
            sum_para = doc.add_paragraph()
            sum_para.add_run(p["summary"])
            sum_para.paragraph_format.left_indent = Pt(12)

            # 分隔线
            border_para = doc.add_paragraph()
            border_para.paragraph_format.space_before = Pt(0)
            border_para.paragraph_format.space_after = Pt(4)
            p_elem = border_para._p
            pPr = p_elem.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "DDDDDD")
            pBdr.append(bottom)
            pPr.append(pBdr)

    add_section("一、过去 24 小时关键政策", report_data.get("policies_24h", []))
    doc.add_paragraph()
    add_section("二、过去一个月主要政策动向", report_data.get("policies_1month", []))
    doc.add_paragraph()
    add_section("三、今年以来主要政策", report_data.get("policies_year", []))
    doc.add_paragraph()

    # ── 趋势 ──
    h = doc.add_heading("四、未来趋势判断", level=1)
    h.runs[0].font.color.rgb = RGBColor(91, 109, 228)
    h.runs[0].font.size = Pt(14)

    nums = ["①", "②", "③", "④", "⑤", "⑥"]
    for i, trend in enumerate(report_data.get("trends", [])):
        tp = doc.add_paragraph()
        nr = tp.add_run(f"{nums[i] if i < len(nums) else '·'}  ")
        nr.font.bold = True
        nr.font.color.rgb = RGBColor(91, 109, 228)
        tp.add_run(trend)

    # ── 页脚 ──
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = footer_p.add_run(
        f"本报告由 AI 自动生成，仅供参考，请以官方发布为准\n"
        f"https://neilbritain.github.io/daily-policy-report"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(180, 180, 180)

    # 保存
    for path in [report_dir / "latest.docx", report_dir / f"report-{date}.docx"]:
        doc.save(path)
    print(f"[OK] Word 已保存: {date}")


def main():
    try:
        data = get_report_data()
        save_word_report(data)
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        raise SystemExit(1)
    except ImportError:
        print("[ERROR] 请安装依赖：pip install python-docx")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
